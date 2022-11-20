"""
作者：zwq
日期:2022年03月13日
"""
import requests
from lxml import etree
# 导入gui库
from tkinter import *
from tkinter import messagebox

# 导入操作docx库
from docx import Document

# 新建空白文档
doc1 = Document()

headers = {
    'Connection': 'keep-alive',
    'Cache-Control': 'max-age=0',
    'Upgrade-Insecure-Requests': '1',
    'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/99.0.4844.51 Safari/537.36',
    'Accept': 'text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8,application/signed-exchange;v=b3;q=0.9',
    'Referer': 'http://m.bd.bendibao.com/news/yqdengji/',
    'Accept-Language': 'zh-CN,zh;q=0.9,en;q=0.8',
}

res = requests.get('http://m.bd.bendibao.com/news/gelizhengce/fengxianmingdan.php', headers=headers, verify=False)
tree = etree.HTML(res.text)

# 主标题
main_title = tree.xpath("//div[@class='border-title']/p[1]/text()")[0]
# 获取时间
mian_time = tree.xpath("//div[@class='city-time-wrap']/p[@class='time']/text()")[0]

# 为文档添加标题
doc1.add_heading(main_title, 0)
# 添加时间
doc1.add_heading(mian_time, 1)


# 高风险地区名单
def high():
    # 标题
    high_title = tree.xpath("//div[@class='fx-item high-fx']/div[2]/text()")[0]
    # 总数量
    high_num = tree.xpath("//div[@class='fx-item high-fx']/div[1]/span/text()")[0]

    # 处理后的标题
    bg_title = high_title + "--------" + high_num + "个"

    # 地区
    diqus = tree.xpath("//div[@class='height info-item']//div[@class='flex-between']/p")
    # 地区数量
    diqu_nums = tree.xpath("//div[@class='height info-item']//div[@class='flex-between']/span[@class='qu-num']/text()")
    # 具体地方
    dfs = tree.xpath("//div[@class='height info-item']/div[@class='info-list']/ul[@class='info-detail']")

    # 设置表格名
    doc1.add_heading(bg_title, 2)
    # 设置表格头
    table = doc1.add_table(rows=1, cols=3)
    th = table.rows[0].cells
    th[0].text = '地区'
    th[1].text = '数量'
    th[2].text = '详细地址'
    # 填入内容

    for diqu, dq_num, df in zip(diqus, diqu_nums, dfs):
        row = table.add_row().cells
        dq = "".join(diqu.xpath(".//span/text()"))
        df = "\n".join(df.xpath(".//span/text()"))
        row[0].text = dq
        row[1].text = dq_num
        row[2].text = "*" + df


# 获取中风险地区
def middle():
    # 标题
    middle_title = tree.xpath("//div[@class='fx-item middle-fx']/div[2]/text()")[0]
    # 总数量
    middle_num = tree.xpath("///div[@class='fx-item middle-fx']/div[1]/span/text()")[0]

    # 处理后的标题
    bg_title1 = middle_title + "--------" + middle_num + "个"

    # 地区
    diqus = tree.xpath("//div[@class='middle info-item']//div[@class='flex-between']/p")
    # 地区数量
    diqu_nums = tree.xpath(
        "//div[@class='middle info-item']//div[@class='flex-between']/span[@class='qu-num']/text()")
    # 具体地方
    dfs = tree.xpath("//div[@class='middle info-item']/div[@class='info-list']/ul[@class='info-detail']")

    # 设置表格名
    doc1.add_heading(bg_title1, 2)
    # 设置表格头
    table = doc1.add_table(rows=1, cols=3)
    th = table.rows[0].cells
    th[0].text = '地区'
    th[1].text = '数量'
    th[2].text = '详细地址'
    # 填入内容

    for diqu1, dq_num1, df1 in zip(diqus, diqu_nums, dfs):
        row = table.add_row().cells
        dq1 = "".join(diqu1.xpath(".//span/text()"))
        df1 = "\n".join(df1.xpath(".//span/text()"))
        row[0].text = dq1
        row[1].text = dq_num1
        row[2].text = "*" + df1


# gui界面
def gui_py():
    wd = Tk()
    wd.title("爬取全国疫情中高风险地区名单")
    wd.geometry('550x400+400+150')
    lbl = Label(wd, text="(请点击按钮进行爬取,下载文件会自动保存到根目录下down_files里)")
    lbl.place(relx=0.15, rely=0.05)
    try:

        btn = Button(wd, text="点击爬取", command=main, height=2, width=10)
        btn.place(relx=0.43, rely=0.35)
    except:
        messagebox.showerror("提示", "发生未知错误,请稍后重试!!!")

    # 调用mainloop()显示主窗口
    wd.mainloop()


# 主函数
def main():
    high()
    middle()
    # 文件名
    docx_time = mian_time[2:12]
    doc1.save(f"{main_title}---{docx_time}.docx")
    messagebox.showinfo("提示", "下载完成")


if __name__ == "__main__":
    # 调用gui函数
    gui_py()
